from PySide6.QtWidgets import (
    QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QListWidget, QListWidgetItem, QLabel,
    QMessageBox, QFileDialog, QDialog
)
from PySide6.QtCore import Qt
from docx import Document

from database import Database
from dialogs import MistakeDialog

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("错题本程序")
        self.resize(500, 600)  # 调整了窗口大小以适应单列表界面
        
        self.db = Database()
        self.init_ui()
        self.load_data()

    def init_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QVBoxLayout(central_widget)
        
        # 顶部操作按钮
        top_btn_layout = QHBoxLayout()
        self.add_btn = QPushButton("添加错题")
        self.add_btn.clicked.connect(self.add_mistake)
        
        self.delete_btn = QPushButton("删除错题")
        self.delete_btn.clicked.connect(self.delete_mistake)
        self.delete_btn.setEnabled(False) # 默认禁用，选中后启用
        
        top_btn_layout.addWidget(self.add_btn)
        top_btn_layout.addWidget(self.delete_btn)
        top_btn_layout.addStretch()
        
        self.list_widget = QListWidget()
        self.list_widget.itemSelectionChanged.connect(self.on_selection_changed)
        self.list_widget.itemDoubleClicked.connect(self.on_item_double_clicked)
        
        self.export_btn = QPushButton("导出为Word")
        self.export_btn.clicked.connect(self.export_to_word)
        
        main_layout.addLayout(top_btn_layout)
        main_layout.addWidget(QLabel("错题列表 (双击查看/编辑详情)"))
        main_layout.addWidget(self.list_widget)
        main_layout.addWidget(self.export_btn)

    def load_data(self):
        self.list_widget.clear()
        mistakes = self.db.get_all_mistakes()
        for mistake in mistakes:
            item = QListWidgetItem(mistake[1])
            item.setData(Qt.UserRole, mistake)
            self.list_widget.addItem(item)
            
        self.on_selection_changed()

    def on_selection_changed(self):
        has_selection = len(self.list_widget.selectedItems()) > 0
        self.delete_btn.setEnabled(has_selection)

    def on_item_double_clicked(self, item):
        mistake_data = item.data(Qt.UserRole)
        # 以查看模式打开，内部可以切换到编辑模式
        dialog = MistakeDialog(self, mistake_data, is_view_mode=True)
        if dialog.exec() == QDialog.Accepted:
            title, question, answer, subject = dialog.get_data()
            if not title or not question or not answer:
                QMessageBox.warning(self, "错误", "所有字段（包括答案）都不能为空！")
                return
            # 如果接受了，说明进行了编辑保存
            self.db.update_mistake(mistake_data[0], title, question, answer, subject)
            self.load_data()

    def add_mistake(self):
        dialog = MistakeDialog(self, is_view_mode=False)
        if dialog.exec() == QDialog.Accepted:
            title, question, answer, subject = dialog.get_data()
            if not title or not question or not answer:
                QMessageBox.warning(self, "错误", "所有字段（包括答案）都不能为空！")
                return
            self.db.add_mistake(title, question, answer, subject)
            self.load_data()

    def delete_mistake(self):
        selected_items = self.list_widget.selectedItems()
        if not selected_items:
            return
            
        reply = QMessageBox.question(self, '确认删除', '确定要删除这道错题吗？',
                                     QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        
        if reply == QMessageBox.Yes:
            mistake_data = selected_items[0].data(Qt.UserRole)
            self.db.delete_mistake(mistake_data[0])
            self.load_data()

    def export_to_word(self):
        mistakes = self.db.get_all_mistakes()
        if not mistakes:
            QMessageBox.information(self, "提示", "当前没有错题可以导出！")
            return
            
        file_path, _ = QFileDialog.getSaveFileName(self, "导出Word文档", "我的错题本", "Word Documents (*.docx)")
        if not file_path:
            return
            
        try:
            doc = Document()
            doc.add_heading('我的错题本', 0)
            
            for index, mistake in enumerate(mistakes, 1):
                # 只输出序号和题干内容，不再输出标题
                p = doc.add_paragraph()
                run = p.add_run(f"{index}. ")
                run.bold = True
                
                # 直接将题干内容接在序号后面，如果题干有换行，保持换行
                p.add_run(mistake[2])
                
                # 添加一个空行分隔不同的题目
                doc.add_paragraph()
                
            doc.save(file_path)
            QMessageBox.information(self, "成功", f"错题本已成功导出至：\n{file_path}")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出失败：{str(e)}")
